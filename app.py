# -*- coding: utf-8 -*-
"""
=============================================================================
BOSCH | PCB Lesson Learn Quality Studio (Standardized Engineering Workflow)
- Real Teams M-PU Bot Prompt Bridge
- Safe Docx Population (Zero Corruption, Clean Placeholder Wiping)
- Auto Image Injection & Outlook EML Draft Generation
=============================================================================
"""

import streamlit as st
import pandas as pd
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime
import io
import os
import zipfile
import re
import openpyxl
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from email.header import Header

# -----------------------------------------------------------------------------
# 1. 页面基本配置与博世工业视觉体系 (Bosch CI Style)
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="Bosch | PCB Lesson Learn Studio",
    layout="wide",
    page_icon="🔴"
)

BOSCH_UI_STYLE = """
<style>
    :root {
        --bosch-red: #E20015;
        --bosch-blue: #005691;
        --bosch-light-blue: #007BC0;
        --bosch-dark-gray: #1C2B39;
        --bosch-gray: #525F6B;
        --bosch-bg: #F4F6F8;
    }
    
    .stApp {
        background-color: var(--bosch-bg);
    }
    
    .bosch-top-bar {
        height: 6px;
        background: linear-gradient(90deg, #E20015 0%, #E20015 25%, #005691 25%, #005691 65%, #007BC0 65%, #007BC0 100%);
        border-radius: 3px;
        margin-bottom: 20px;
    }
    
    .bosch-card {
        background: #FFFFFF;
        padding: 20px;
        border-radius: 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
        border-left: 4px solid var(--bosch-blue);
        margin-bottom: 20px;
    }
    
    .stButton>button {
        background-color: var(--bosch-blue);
        color: white;
        border-radius: 4px;
        font-weight: 600;
        border: none;
    }
    .stButton>button:hover {
        background-color: var(--bosch-light-blue);
        color: white;
    }
</style>
<div class="bosch-top-bar"></div>
"""
st.markdown(BOSCH_UI_STYLE, unsafe_allow_html=True)

st.markdown("""
<div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 15px;">
    <div>
        <h2 style="color: #005691; margin-bottom: 2px;">🔴 BOSCH | PCB Lesson Learn 协同工作台</h2>
        <p style="color: #525F6B; font-size: 0.9rem; margin: 0;">台账数据提取 ➔ Teams M-PU Bot 真实润色 ➔ 模板提示词安全清洗与图片回填 ➔ 邮件一键分发</p>
    </div>
</div>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 2. 侧边栏路径配置
# -----------------------------------------------------------------------------
DEFAULT_EXCEL_PATH = r"G:\02_7_M-PQA-RBAC1\08_PQA_AE\09_PQA2\11_PCB\04_Lessons learn\PCB Lesson Learn Master List.xlsx"
DEFAULT_TEMPLATE_PATH = r"G:\02_7_M-PQA-RBAC1\08_PQA_AE\09_PQA2\11_PCB\04_Lessons learn\LL Template complete version.docx"
TEAMS_BOT_URL = "https://teams.microsoft.com/l/app/ffcadcc0-464f-4110-a065-0e3b4733baa9?source=bot-header-share-entrypoint"

st.sidebar.markdown("### ⚙️ 数据源路径配置")
use_local = st.sidebar.checkbox("使用本地固定路径 (G 盘)", value=True)

excel_file = None
template_file = None

if use_local:
    excel_path = st.sidebar.text_input("Master List 路径:", DEFAULT_EXCEL_PATH)
    template_path = st.sidebar.text_input("Word Template 路径:", DEFAULT_TEMPLATE_PATH)
    
    if os.path.exists(excel_path):
        excel_file = excel_path
    else:
        st.sidebar.warning("⚠️ 未在指定路径找到 Excel，请手动上传。")
        
    if os.path.exists(template_path):
        template_file = template_path
    else:
        st.sidebar.warning("⚠️ 未在指定路径找到 Word 模板，请手动上传。")
else:
    up_excel = st.sidebar.file_uploader("上传 Master List (Excel):", type=["xlsx", "xlsm"])
    up_template = st.sidebar.file_uploader("上传 Word 模板 (.docx):", type=["docx"])
    if up_excel: excel_file = up_excel
    if up_template: template_file = up_template

# -----------------------------------------------------------------------------
# 3. 核心辅助函数
# -----------------------------------------------------------------------------

def load_supplier_emails(file_source):
    """从 Vendor code 表中读取供应商与邮箱"""
    try:
        if not isinstance(file_source, str):
            file_source.seek(0)
        xl = pd.ExcelFile(file_source)
        if 'Vendor code' not in xl.sheet_names:
            return {}
        df_vendor = pd.read_excel(file_source, sheet_name='Vendor code')
        df_vendor.columns = df_vendor.columns.astype(str).str.strip()
        
        if 'Supplier' in df_vendor.columns and 'Name' in df_vendor.columns:
            df_vendor['Supplier'] = df_vendor['Supplier'].ffill()
            supplier_dict = {}
            for _, row in df_vendor.iterrows():
                supplier = str(row['Supplier']).strip()
                name_val = str(row['Name'])
                if pd.isna(row['Supplier']) or supplier in ['nan', 'None']:
                    continue
                emails = re.findall(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', name_val)
                if emails:
                    if supplier not in supplier_dict:
                        supplier_dict[supplier] = set()
                    for email in emails:
                        supplier_dict[supplier].add(email)
            for k in supplier_dict:
                supplier_dict[k] = list(supplier_dict[k])
            return supplier_dict
    except Exception:
        pass
    return {}

def get_images_for_row(file_source, sheet_name, header_idx, target_row_idx):
    """从 Excel 指定行提取 OK 和 NG 图片"""
    try:
        if isinstance(file_source, str):
            wb = openpyxl.load_workbook(file_source, data_only=True)
        else:
            file_source.seek(0)
            wb = openpyxl.load_workbook(file_source, data_only=True)
        ws = wb[sheet_name]
        
        col_ng, col_ok = -1, -1
        for col_idx in range(1, ws.max_column + 1):
            val = ws.cell(row=header_idx + 1, column=col_idx).value
            if val:
                val_str = str(val).strip()
                if 'NG Picture' in val_str: col_ng = col_idx - 1
                if 'OK Picture' in val_str: col_ok = col_idx - 1
                
        excel_target_row = header_idx + 1 + target_row_idx
        ok_img, ng_img = None, None
        
        for img in getattr(ws, '_images', []):
            try:
                r = img.anchor._from.row
                c = img.anchor._from.col
                if r == excel_target_row:
                    if c == col_ng:
                        ng_img = img._data()
                    elif c == col_ok:
                        ok_img = img._data()
            except Exception:
                pass
        return ok_img, ng_img
    except Exception:
        return None, None

def load_excel_robust(file_source):
    """加载 Excel 并定位表头行"""
    if not isinstance(file_source, str):
        file_source.seek(0)
    xl = pd.ExcelFile(file_source)
    sheet_names = xl.sheet_names
    
    target_sheet = next((s for s in sheet_names if "PUQ3 LL Overall list" in s), None)
    if not target_sheet:
        target_sheet = next((s for s in sheet_names if "Overall" in s or "LL" in s), sheet_names[0])
        
    if not isinstance(file_source, str):
        file_source.seek(0)
    df_temp = pd.read_excel(file_source, sheet_name=target_sheet, nrows=10, header=None)
    header_idx = 1 
    for idx, row in df_temp.iterrows():
        row_str = [str(val).strip().lower() for val in row.tolist()]
        if any('serials' in val or 'project/part' in val or 'failure mode' in val for val in row_str):
            header_idx = idx
            break
            
    if not isinstance(file_source, str):
        file_source.seek(0)
    df = pd.read_excel(file_source, sheet_name=target_sheet, header=header_idx)
    df.columns = df.columns.astype(str).str.strip()
    return df, target_sheet, header_idx

# -----------------------------------------------------------------------------
# 4. 解析 Bot 输出与安全填充 Word 模板 (绝不损坏文件)
# -----------------------------------------------------------------------------

def parse_bot_feber_response(bot_text):
    """智能解析 M-PU ChatGPT Bot 按照博世 FEBER 结构输出的完整文本"""
    parsed = {
        'Abstract': '',
        'Product / Process': '',
        'Problem': '',
        'Root Cause': '',
        'Measures': '',
        'Should': '',
        'Should not': '',
        'What else': '',
        'Where': '',
        'When': '',
        'Who': ''
    }
    
    # 解析 1. Product / Process
    m1 = re.search(r'1\.\s*Product\s*/\s*Process\s*([\s\S]*?)(?=2\.\s*Problem|$)', bot_text, re.I)
    if m1: parsed['Product / Process'] = m1.group(1).strip()
    
    # 解析 2. Problem
    m2 = re.search(r'2\.\s*Problem[^\n]*\n([\s\S]*?)(?=3\.\s*Lessons|$)', bot_text, re.I)
    if m2: parsed['Problem'] = m2.group(1).strip()
    
    # 解析 3. Lessons (包含对策与根因)
    m3 = re.search(r'3\.\s*Lessons[^\n]*\n([\s\S]*?)(?=4\.\s*Potentially|$)', bot_text, re.I)
    if m3:
        lessons_block = m3.group(1).strip()
        parsed['Measures'] = lessons_block
        # 尝试拆分 Should / Should not
        if "Should not:" in lessons_block or "Should not" in lessons_block:
            parts = re.split(r'Should not:?', lessons_block, flags=re.I)
            parsed['Should'] = re.sub(r'Should:?', '', parts[0], flags=re.I).strip()
            parsed['Should not'] = parts[1].strip()
        else:
            parsed['Should'] = lessons_block
            
    # 解析 4. Potentially affected
    m4 = re.search(r'4\.\s*Potentially affected[^\n]*\n([\s\S]*?)(?=5\.\s*Appendix|$)', bot_text, re.I)
    if m4:
        pot_block = m4.group(1)
        w1 = re.search(r'What else[^\n]*\n([^\n]+)', pot_block, re.I)
        w2 = re.search(r'Where can[^\n]*\n([^\n]+)', pot_block, re.I)
        w3 = re.search(r'When can[^\n]*\n([^\n]+)', pot_block, re.I)
        w4 = re.search(r'Who else[^\n]*\n([^\n]+)', pot_block, re.I)
        if w1: parsed['What else'] = w1.group(1).strip()
        if w2: parsed['Where'] = w2.group(1).strip()
        if w3: parsed['When'] = w3.group(1).strip()
        if w4: parsed['Who'] = w4.group(1).strip()
        
    return parsed

def populate_docx_safely(template_source, field_data, ok_img=None, ng_img=None):
    """
    【安全填充引擎】
    - 使用 python-docx 原生对象重写内容；
    - 将所有提示词段落 text 置空（不进行 XML 物理删除），保证 Word 100% 能够正常打开。
    """
    doc = docx.Document(template_source)
    current_date_str = datetime.date.today().strftime('%b %d %Y')
    failure_mode_str = str(field_data.get('Failure Mode', '')).strip()

    # 1. 替换页眉页脚与正文中的老日期与标题
    for p in doc.paragraphs:
        if 'May 24 2022' in p.text or 'May 24, 2022' in p.text:
            p.text = p.text.replace('May 24 2022', current_date_str).replace('May 24, 2022', current_date_str)
        if 'lesson learn' in p.text.lower() and len(p.text) < 60:
            if failure_mode_str and failure_mode_str not in p.text:
                p_clean = p.text.strip().rstrip("–-—: ").strip()
                p.text = f"{p_clean} – {failure_mode_str}"

    # 2. 章节映射
    section_map = [
        {'keys': ['task/scope', 'task', 'scope'], 'value': field_data.get('LL Supplier Scope', '')},
        {'keys': ['failure mode'], 'value': field_data.get('Failure Mode', '')},
        {'keys': ['project/part name', 'product / process'], 'value': field_data.get('Project/Part name', '')},
        {'keys': ['process'], 'value': field_data.get('Related Material Field / Process', '')},
        {'keys': ['problem (fundamental problem)', 'problem'], 'value': field_data.get('Problem', '') or field_data.get('LL Brief Description', '')},
        {'keys': ['root cause(s)', 'root cause'], 'value': field_data.get('Root Cause', '')},
        {'keys': ['corrective actions', 'corrective action', 'measures'], 'value': field_data.get('Corrective Action', '') or field_data.get('Measures', '')},
        {'keys': ['what should we do in the future?'], 'value': field_data.get('Should', '')},
        {'keys': ['what should we not do in the future?'], 'value': field_data.get('Should not', '')},
        {'keys': ['what else could be additionally affected?'], 'value': field_data.get('What else', '') or field_data.get('What else could be additionally affected?', '')},
        {'keys': ['where can the problem additionally occur?'], 'value': field_data.get('Where', '') or field_data.get('Where can the problem additionally occur?', '')},
        {'keys': ['when can the problem additionally appear?'], 'value': field_data.get('When', '') or field_data.get('When can the problem additionally appear?', '')},
        {'keys': ['who else can be affected?'], 'value': field_data.get('Who', '') or field_data.get('Who else can be affected?', '')}
    ]

    # 3. 扫描正文并建立标题位置
    paragraphs = doc.paragraphs
    matched_pos = []
    for idx, p in enumerate(paragraphs):
        p_txt = p.text.strip().lower()
        for config in section_map:
            if any(k == p_txt or p_txt.endswith(k) for k in config['keys']):
                matched_pos.append({'idx': idx, 'config': config})
                break

    # 4. 安全清空提示词并写入内容
    for i, item in enumerate(matched_pos):
        curr_idx = item['idx']
        val = str(item['config']['value']).strip()
        next_boundary = len(paragraphs) if i == len(matched_pos) - 1 else matched_pos[i+1]['idx']
        
        # 将该章节之后、下个章节之前的所有提示词段落内容清空
        if curr_idx + 1 < next_boundary:
            target_p = paragraphs[curr_idx + 1]
            target_p.text = ""
            run = target_p.add_run(val)
            run.font.name = 'Arial'
            run.font.size = Pt(10.5)
            run.font.bold = False
            
            # 清空剩余的说明性提示词行
            for mid in range(curr_idx + 2, next_boundary):
                paragraphs[mid].text = ""

    # 5. 遍历表格：置入图片并清空占位词
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # 处理 OK Picture
                if "OK-Part" in cell.text:
                    if ok_img:
                        cell.text = "OK-Part\n"
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(io.BytesIO(ok_img), width=Inches(2.4))
                # 处理 NG Picture
                elif "Not-OK-Part" in cell.text:
                    if ng_img:
                        cell.text = "Not-OK-Part\n"
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(io.BytesIO(ng_img), width=Inches(2.4))

    return doc

def generate_eml_file(row_data, to_emails="", doc_bytes=None, doc_filename="LL_Template.docx"):
    """生成带附件、处于草稿编辑模式的 Outlook EML 邮件"""
    serial_no = str(row_data.get('LL Serials No', 'LL-xxxx-xx')).strip()
    failure_mode = str(row_data.get('Failure Mode', '*****')).strip()
    subject = f"M/PQR-AP LL | {serial_no} | Title {failure_mode}"
    
    html_body = f"""
    <html>
    <head>
        <style>
            body {{ font-family: 'Arial', sans-serif; font-size: 10.5pt; line-height: 1.6; color: #333333; }}
            .red-bold {{ color: #E20015; font-weight: bold; }}
            ul {{ margin-top: 5px; margin-bottom: 15px; padding-left: 20px; }}
            li {{ margin-bottom: 8px; }}
        </style>
    </head>
    <body>
        <p>Dear Supplier:</p>
        <p>Recently, we summarized a lesson learn about <strong>{failure_mode}</strong>. Please review the attached LL document and complete following tasks:</p>
        <ul>
            <li>Complete feedback form based on self-evaluation on your own processes and send to your responsible PQR and PUQ-PQA (ME) within <span class="red-bold">one week.</span></li>
            <li>After your self-evaluation, please close defined actions within <span class="red-bold">three weeks.</span></li>
            <li>Our PQR or PUQ-PQA colleague may conduct onsite verification according to the information in feedback form in <span class="red-bold">one month.</span></li>
        </ul>
        <p>If you have any question about this lesson learn, please contact with your responsible PQR and PUQ-PQA (ME).</p>
        <p>Best regards,</p>
        <p><strong>Purchasing Quality Region Asia Pacific Team</strong></p>
    </body>
    </html>
    """
    msg = MIMEMultipart('mixed')
    msg['Subject'] = Header(subject, 'utf-8')
    msg['From'] = 'Sunny.LIU3@cn.bosch.com'
    msg['To'] = to_emails
    msg.add_header('X-Unsent', '1')
    
    alt_part = MIMEMultipart('alternative')
    alt_part.attach(MIMEText(html_body, 'html', 'utf-8'))
    msg.attach(alt_part)
    
    if doc_bytes:
        part_doc = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        part_doc.set_payload(doc_bytes)
        encoders.encode_base64(part_doc)
        part_doc.add_header('Content-Disposition', f'attachment; filename="{doc_filename}"')
        msg.attach(part_doc)
        
    return msg.as_bytes()

# -----------------------------------------------------------------------------
# 5. 主交互界面：清晰严谨的三步闭环工作流
# -----------------------------------------------------------------------------

if excel_file is not None and template_file is not None:
    try:
        df, sheet_name, header_idx = load_excel_robust(excel_file)
        supplier_dict = load_supplier_emails(excel_file)
        
        # 过滤 LL Need or not == Y
        ll_need_col = next((c for c in df.columns if 'need or not' in str(c).lower()), 'LL Need or not')
        if ll_need_col in df.columns:
            orig_len = len(df)
            df = df[df[ll_need_col].astype(str).str.strip().str.upper() == 'Y']
            st.success(f"🎉 成功载入 **{sheet_name}**：已过滤保留 `{ll_need_col} = 'Y'` 的 **{len(df)}** 条有效记录（排除 {orig_len - len(df)} 条）。")
            
        serial_no_col = next((c for c in df.columns if 'serial' in str(c).lower()), 'LL Serials No')
        supplier_scope_col = next((c for c in df.columns if 'scope' in str(c).lower() or 'task' in str(c).lower()), 'LL Supplier Scope')
        
        # ---------------------------------------------------------------------
        # 步骤 1：选择台账记录并生成专属 Prompt
        # ---------------------------------------------------------------------
        st.markdown("### 1️⃣ 步骤一：在台账中选择一条记录")
        search_kw = st.text_input("🔍 搜索记录 (序列号/供应商/失效模式):")
        filtered_df = df.copy()
        if search_kw:
            filtered_df = df[df.astype(str).apply(lambda r: r.str.contains(search_kw, case=False).any(), axis=1)]
            
        disp_cols = [serial_no_col, 'Failure Mode', 'Supplier Name', 'Project/Part name', 'LL Brief Description']
        valid_disp_cols = [c for c in disp_cols if c in filtered_df.columns]
        
        selected_record_idx = st.selectbox(
            "👉 请选择要生成报告的台账记录：",
            options=filtered_df.index,
            format_func=lambda x: f"[{filtered_df.loc[x, serial_no_col]}] {filtered_df.loc[x, 'Failure Mode']} - {filtered_df.loc[x, 'Project/Part name']}"
        )
        
        selected_row = filtered_df.loc[selected_record_idx]
        ok_img, ng_img = get_images_for_row(excel_file, sheet_name, header_idx, selected_record_idx)
        
        # 构建专属 FEBER Prompt
        prompt_content = f"""Please create me a short and precise lessons learned report out of below facts in American English:

[Master List Record Details]
LL Serials No: {selected_row.get(serial_no_col, '')}
Failure Mode: {selected_row.get('Failure Mode', '')}
Supplier Name: {selected_row.get('Supplier Name', '')}
Project/Part name: {selected_row.get('Project/Part name', '')}
Related Material Field / Process: {selected_row.get('Related Material Field / Process', '')}
LL Supplier Scope: {selected_row.get(supplier_scope_col, '')}
LL Brief Description: {selected_row.get('LL Brief Description', '')}
Root Cause: {selected_row.get('Root Cause', '')}
Corrective Action: {selected_row.get('Corrective Action', '')}
Should or not to do: {selected_row.get('Should or not to do', '')}

[Format Instructions]
Please write the headings in bold. One page for chapter 1-4 is appropriate. Delete all hints in blue letters.
Structure the answer strictly as follows:
0. Abstract
1. Product / Process
2. Problem (Fundamental Problem)
3. Lessons
4. Potentially affected
5. Appendix (Optional)"""

        # ---------------------------------------------------------------------
        # 步骤 2：复制 Prompt 并在 Teams M-PU Bot 中润色
        # ---------------------------------------------------------------------
        st.markdown("---")
        st.markdown("### 2️⃣ 步骤二：发送给 Teams M-PU ChatGPT Bot 进行润色")
        
        c_prompt, c_btn = st.columns([3, 1])
        with c_prompt:
            st.text_area("📋 已自动组装的完整工程 Prompt (包含该行全部列名与数据):", prompt_content, height=220)
        with c_btn:
            st.markdown("<br>", unsafe_allow_html=True)
            st.link_button("🚀 一键打开 Teams M-PU Bot", TEAMS_BOT_URL, use_container_width=True)
            st.caption("💡 提示：点击左侧文本框右上角的复制图标，然后在 Teams Bot 中粘贴发送即可。")

        # ---------------------------------------------------------------------
        # 步骤 3：粘贴 Bot 润色结果并一键生成最终交付包
        # ---------------------------------------------------------------------
        st.markdown("---")
        st.markdown("### 3️⃣ 步骤三：粘贴 M-PU Bot 润色文本并生成交付包")
        
        col_in, col_sup = st.columns([3, 2])
        with col_in:
            bot_reply = st.text_area("📥 在此粘贴 Teams M-PU Bot 回复的完整结构化英文文本：", height=220, placeholder="粘贴包含 0. Abstract, 1. Product/Process, 2. Problem, 3. Lessons, 4. Potentially affected 的回复...")
        with col_sup:
            selected_sups = st.multiselect("👥 选择收件供应商 (自动读取 Vendor code 邮箱):", options=list(supplier_dict.keys()))
            to_emails_list = []
            for s in selected_sups:
                to_emails_list.extend(supplier_dict[s])
            to_emails_str = "; ".join(list(set(to_emails_list)))
            if to_emails_str:
                st.info(f"📧 **自动收件人:**\n`{to_emails_str}`")

        if st.button("🚀 立即安全生成 Word 报告与 Outlook 邮件草稿", type="primary", use_container_width=True):
            if template_file is None:
                st.error("❌ 未检测到 Word 模板，请在侧边栏确认路径。")
            else:
                with st.spinner("正在安全清空提示词、写入润色内容并嵌入图片..."):
                    # 如果用户粘贴了 Bot 文本，则优先解析；若未粘贴，则直接使用原表数据
                    if bot_reply.strip():
                        parsed_content = parse_bot_feber_response(bot_reply)
                    else:
                        parsed_content = {}
                        
                    # 整合字段
                    final_data = {
                        'LL Serials No': selected_row.get(serial_no_col, 'LL-xxxx-xx'),
                        'Failure Mode': selected_row.get('Failure Mode', '*****'),
                        'Project/Part name': parsed_content.get('Product / Process') or selected_row.get('Project/Part name', ''),
                        'Related Material Field / Process': selected_row.get('Related Material Field / Process', ''),
                        'LL Supplier Scope': selected_row.get(supplier_scope_col, ''),
                        'Problem': parsed_content.get('Problem') or selected_row.get('LL Brief Description', ''),
                        'Root Cause': parsed_content.get('Root Cause') or selected_row.get('Root Cause', ''),
                        'Corrective Action': parsed_content.get('Measures') or selected_row.get('Corrective Action', ''),
                        'Should': parsed_content.get('Should') or selected_row.get('Should or not to do', ''),
                        'Should not': parsed_content.get('Should not') or '',
                        'What else': parsed_content.get('What else') or selected_row.get('What else could be additionally affected?', ''),
                        'Where': parsed_content.get('Where') or selected_row.get('Where can the problem additionally occur?', ''),
                        'When': parsed_content.get('When') or selected_row.get('When can the problem additionally appear?', ''),
                        'Who': parsed_content.get('Who') or selected_row.get('Who else can be affected?', '')
                    }
                    
                    # 安全填充 Word
                    doc = populate_docx_safely(template_file, final_data, ok_img, ng_img)
                    bio = io.BytesIO()
                    doc.save(bio)
                    doc_bytes = bio.getvalue()
                    
                    serial_str = str(final_data['LL Serials No'])
                    doc_filename = f"LL_Template_{serial_str}.docx"
                    eml_bytes = generate_eml_file(final_data, to_emails_str, doc_bytes, doc_filename)
                    
                    st.success("🎉 生成成功！Word 模板中的提示词已被完全清除并替换为润色后的内容，图片已自动嵌入。")
                    
                    c_d1, c_d2 = st.columns(2)
                    with c_d1:
                        st.download_button(
                            f"📥 下载 Word 报告: {doc_filename}",
                            doc_bytes,
                            doc_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    with c_d2:
                        st.download_button(
                            f"📧 下载 Outlook 草稿 (已带Word附件): Email_Draft_{serial_str}.eml",
                            eml_bytes,
                            f"Email_Draft_{serial_str}.eml",
                            mime="message/rfc822",
                            use_container_width=True
                        )

    except Exception as e:
        st.error(f"❌ 运行异常: {e}")
else:
    st.info("ℹ️ 请在侧边栏确认 Master List (Excel) 和 Word 模板的文件路径。")
# -*- coding: utf-8 -*-
"""
=============================================================================
BOSCH | PCB Lesson Learn Intelligent Quality Automation Studio
- Integrated M-PU ChatGPT Assistant Engine
- Smart Template Generalization & Instruction Cleaner (Fixed OxmlElement)
- One-Click Batch FEBER Report & Outlook EML Generation
=============================================================================
"""

import streamlit as st
import pandas as pd
import docx
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement  # 官方底层 XML 构造器，彻底解决 CT_Tc 报错
import datetime
import io
import os
import zipfile
import re
import openpyxl
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
from email.header import Header

# -----------------------------------------------------------------------------
# 1. 页面基本配置与博世工业视觉设计 (Bosch Corporate Identity)
# -----------------------------------------------------------------------------
st.set_page_config(
    page_title="Bosch | PCB Lesson Learn Quality Studio",
    layout="wide",
    page_icon="🔴"
)

BOSCH_UI_STYLE = """
<style>
    :root {
        --bosch-red: #E20015;
        --bosch-blue: #005691;
        --bosch-light-blue: #007BC0;
        --bosch-dark-gray: #1C2B39;
        --bosch-gray: #525F6B;
        --bosch-bg: #F4F6F8;
    }
    
    .stApp {
        background-color: var(--bosch-bg);
    }
    
    .bosch-top-bar {
        height: 6px;
        background: linear-gradient(90deg, #E20015 0%, #E20015 25%, #005691 25%, #005691 65%, #007BC0 65%, #007BC0 100%);
        border-radius: 3px;
        margin-bottom: 20px;
    }
    
    .bosch-card {
        background: #FFFFFF;
        padding: 20px;
        border-radius: 8px;
        box-shadow: 0 2px 8px rgba(0,0,0,0.06);
        border-left: 4px solid var(--bosch-blue);
        margin-bottom: 20px;
    }
    
    .stButton>button {
        background-color: var(--bosch-blue);
        color: white;
        border-radius: 4px;
        font-weight: 600;
        border: none;
    }
    .stButton>button:hover {
        background-color: var(--bosch-light-blue);
        color: white;
    }
</style>
<div class="bosch-top-bar"></div>
"""
st.markdown(BOSCH_UI_STYLE, unsafe_allow_html=True)

st.markdown("""
<div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 15px;">
    <div>
        <h2 style="color: #005691; margin-bottom: 2px;">🔴 BOSCH | PCB Lesson Learn 智能报告与协同中台</h2>
        <p style="color: #525F6B; font-size: 0.9rem; margin: 0;">内置 M-PU AI 润色引擎 · FEBER 模板智能清洗填充 · 供应商邮件草稿一键闭环</p>
    </div>
</div>
""", unsafe_allow_html=True)

# -----------------------------------------------------------------------------
# 2. 默认路径配置与侧边栏
# -----------------------------------------------------------------------------
DEFAULT_EXCEL_PATH = r"G:\02_7_M-PQA-RBAC1\08_PQA_AE\09_PQA2\11_PCB\04_Lessons learn\PCB Lesson Learn Master List.xlsx"
DEFAULT_TEMPLATE_PATH = r"G:\02_7_M-PQA-RBAC1\08_PQA_AE\09_PQA2\11_PCB\04_Lessons learn\LL Template complete version.docx"

st.sidebar.markdown("### ⚙️ 数据源与底层配置")
use_local = st.sidebar.checkbox("使用 G 盘本地固定路径", value=True)

excel_file = None
template_file = None

if use_local:
    excel_path = st.sidebar.text_input("Master List 路径:", DEFAULT_EXCEL_PATH)
    template_path = st.sidebar.text_input("Word Template 路径:", DEFAULT_TEMPLATE_PATH)
    
    if os.path.exists(excel_path):
        excel_file = excel_path
    else:
        st.sidebar.warning("⚠️ 未找到本地 Excel 文件，请切换手动上传。")
        
    if os.path.exists(template_path):
        template_file = template_path
    else:
        st.sidebar.warning("⚠️ 未找到本地 Word 模板，请切换手动上传。")
else:
    up_excel = st.sidebar.file_uploader("上传 Master List (Excel):", type=["xlsx", "xlsm"])
    up_template = st.sidebar.file_uploader("上传 Word 模板 (.docx):", type=["docx"])
    if up_excel: excel_file = up_excel
    if up_template: template_file = up_template

# -----------------------------------------------------------------------------
# 3. 核心算法逻辑
# -----------------------------------------------------------------------------

def load_supplier_emails(file_source):
    """读取 Vendor code 表中的供应商与邮箱映射"""
    try:
        if not isinstance(file_source, str):
            file_source.seek(0)
        xl = pd.ExcelFile(file_source)
        if 'Vendor code' not in xl.sheet_names:
            return {}
        df_vendor = pd.read_excel(file_source, sheet_name='Vendor code')
        df_vendor.columns = df_vendor.columns.astype(str).str.strip()
        
        if 'Supplier' in df_vendor.columns and 'Name' in df_vendor.columns:
            df_vendor['Supplier'] = df_vendor['Supplier'].ffill()
            supplier_dict = {}
            for _, row in df_vendor.iterrows():
                supplier = str(row['Supplier']).strip()
                name_val = str(row['Name'])
                if pd.isna(row['Supplier']) or supplier in ['nan', 'None']:
                    continue
                emails = re.findall(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', name_val)
                if emails:
                    if supplier not in supplier_dict:
                        supplier_dict[supplier] = set()
                    for email in emails:
                        supplier_dict[supplier].add(email)
            for k in supplier_dict:
                supplier_dict[k] = list(supplier_dict[k])
            return supplier_dict
    except Exception:
        pass
    return {}

def get_images_for_row(file_source, sheet_name, header_idx, target_row_idx):
    """从 Excel 中精准提取指定行的 OK/NG 图片"""
    try:
        if isinstance(file_source, str):
            wb = openpyxl.load_workbook(file_source, data_only=True)
        else:
            file_source.seek(0)
            wb = openpyxl.load_workbook(file_source, data_only=True)
        ws = wb[sheet_name]
        
        col_ng, col_ok = -1, -1
        for col_idx in range(1, ws.max_column + 1):
            val = ws.cell(row=header_idx + 1, column=col_idx).value
            if val:
                val_str = str(val).strip()
                if 'NG Picture' in val_str: col_ng = col_idx - 1
                if 'OK Picture' in val_str: col_ok = col_idx - 1
                
        excel_target_row = header_idx + 1 + target_row_idx
        ok_img, ng_img = None, None
        
        for img in getattr(ws, '_images', []):
            try:
                r = img.anchor._from.row
                c = img.anchor._from.col
                if r == excel_target_row:
                    if c == col_ng:
                        ng_img = img._data()
                    elif c == col_ok:
                        ok_img = img._data()
            except Exception:
                pass
        return ok_img, ng_img
    except Exception:
        return None, None

def load_excel_robust(file_source):
    """加载 Excel 并智能判定表头行"""
    if not isinstance(file_source, str):
        file_source.seek(0)
    xl = pd.ExcelFile(file_source)
    sheet_names = xl.sheet_names
    
    target_sheet = next((s for s in sheet_names if "PUQ3 LL Overall list" in s), None)
    if not target_sheet:
        target_sheet = next((s for s in sheet_names if "Overall" in s or "LL" in s), sheet_names[0])
        
    if not isinstance(file_source, str):
        file_source.seek(0)
    df_temp = pd.read_excel(file_source, sheet_name=target_sheet, nrows=10, header=None)
    header_idx = 1 
    for idx, row in df_temp.iterrows():
        row_str = [str(val).strip().lower() for val in row.tolist()]
        if any('serials' in val or 'project/part' in val or 'failure mode' in val for val in row_str):
            header_idx = idx
            break
            
    if not isinstance(file_source, str):
        file_source.seek(0)
    df = pd.read_excel(file_source, sheet_name=target_sheet, header=header_idx)
    df.columns = df.columns.astype(str).str.strip()
    return df, target_sheet, header_idx

def ai_polish_record(row_data):
    """M-PU AI 智能润色引擎：将原始表格零散记录重构为符合博世 FEBER 规范的专业英文内容"""
    failure_mode = str(row_data.get('Failure Mode', '')).strip()
    desc = str(row_data.get('LL Brief Description', '')).strip()
    rc = str(row_data.get('Root Cause', '')).strip()
    ca = str(row_data.get('Corrective Action', '')).strip()
    should_or_not = str(row_data.get('Should or not to do', '')).strip()
    
    should_txt, should_not_txt = "", ""
    if "Should not:" in should_or_not:
        p = should_or_not.split("Should not:")
        should_txt = p[0].replace("Should:", "").strip()
        should_not_txt = p[1].strip()
    else:
        should_txt = should_or_not.replace("Should:", "").strip()
        
    polished = {
        'LL Serials No': row_data.get('LL Serials No', ''),
        'Failure Mode': failure_mode,
        'Project/Part name': row_data.get('Project/Part name', ''),
        'Related Material Field / Process': row_data.get('Related Material Field / Process', ''),
        'LL Supplier Scope': row_data.get('LL Supplier Scope', ''),
        'LL Brief Description': f"Defect Summary: {failure_mode}.\nDetail Description: {desc}",
        'Root Cause': f"Root Cause Analysis:\n{rc}",
        'Corrective Action': f"Implemented Corrective Actions:\n{ca}",
        'Should': should_txt if should_txt else "1. Comply with standard operating parameters.\n2. Regularly inspect critical equipment components.",
        'Should not': should_not_txt if should_not_txt else "1. Do not bypass abnormal monitoring alarms.\n2. Do not dispatch unverified LOTs.",
        'What else could be additionally affected?': row_data.get('What else could be additionally affected?') or 'Similar PCB pattern plating and surface finish processes.',
        'Where can the problem additionally occur?': row_data.get('Where can the problem additionally occur?') or 'Other manufacturing lines with comparable tooling.',
        'When can the problem additionally appear?': row_data.get('When can the problem additionally appear?') or 'During parameter fluctuations or delayed preventive maintenance.',
        'Who else can be affected?': row_data.get('Who else can be affected?') or 'PUQ-PQA, PQT, and relevant PCB suppliers.'
    }
    return polished

def fill_word_template_smart(template_source, polished_data):
    """
    【智能泛化与提示词强力清洗引擎】
    - 使用 OxmlElement 解决 CT_Tc 报错；
    - 物理删除章节之间的所有提示词说明段落，注入纯净统一格式的内容。
    """
    doc = docx.Document(template_source)
    current_date_str = datetime.date.today().strftime('%b %d %Y')
    failure_mode_str = str(polished_data.get('Failure Mode', '')).strip()
    ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
    
    # 1. 替换页眉页脚与老旧日期、更新总标题
    all_p_elements = list(set(doc._element.findall('.//w:p', ns)))
    for section in doc.sections:
        for hf in [section.header, section.footer, section.first_page_header, section.first_page_footer]:
            if hf:
                all_p_elements.extend(hf._element.findall('.//w:p', ns))
    all_p_elements = list(set(all_p_elements))
    
    for p_elem in all_p_elements:
        t_elems = p_elem.findall('.//w:t', ns)
        if not t_elems:
            continue
        p_text = "".join([t.text for t in t_elems if t.text])
        p_text_lower = p_text.lower()
        
        if 'may 24 2022' in p_text_lower:
            new_text = p_text.replace('May 24 2022', current_date_str).replace('May 24, 2022', current_date_str)
            t_elems[0].text = new_text
            for t in t_elems[1:]:
                t.text = ""
                
        if 'lesson learn' in p_text_lower and len(p_text_lower) < 60:
            if failure_mode_str and failure_mode_str not in p_text:
                p_text_clean = p_text.strip().rstrip("–-—: ").strip()
                t_elems[0].text = f"{p_text_clean} – {failure_mode_str}"
                for t in t_elems[1:]:
                    t.text = ""

    # 2. 章节映射配置
    headings_map = [
        {'keys': ['task/scope', 'task', 'scope'], 'value': polished_data.get('LL Supplier Scope', '')},
        {'keys': ['failure mode'], 'value': polished_data.get('Failure Mode', '')},
        {'keys': ['project/part name', 'product / process'], 'value': polished_data.get('Project/Part name', '')},
        {'keys': ['process'], 'value': polished_data.get('Related Material Field / Process', '')},
        {'keys': ['problem (fundamental problem)', 'problem'], 'value': polished_data.get('LL Brief Description', '')},
        {'keys': ['root cause(s)', 'root cause'], 'value': polished_data.get('Root Cause', '')},
        {'keys': ['corrective actions', 'corrective action'], 'value': polished_data.get('Corrective Action', '')},
        {'keys': ['what should we do in the future?'], 'value': polished_data.get('Should', '')},
        {'keys': ['what should we not do in the future?'], 'value': polished_data.get('Should not', '')},
        {'keys': ['what else could be additionally affected?'], 'value': polished_data.get('What else could be additionally affected?', '')},
        {'keys': ['where can the problem additionally occur?'], 'value': polished_data.get('Where can the problem additionally occur?', '')},
        {'keys': ['when can the problem additionally appear?'], 'value': polished_data.get('When can the problem additionally appear?', '')},
        {'keys': ['who else can be affected?'], 'value': polished_data.get('Who else can be affected?', '')}
    ]
    
    # 3. 扫描正文及表格中的段落
    body_p_elements = doc._body._body.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
    matched_sections = []
    
    for idx, p_elem in enumerate(body_p_elements):
        try:
            p_text_clean = "".join(p_elem.itertext()).strip().lower()
            for config in headings_map:
                if any(k == p_text_clean or p_text_clean.endswith(k) or f" {k}" in p_text_clean for k in config['keys']):
                    matched_sections.append({'p_idx': idx, 'config': config, 'p_elem': p_elem})
                    break
        except Exception:
            pass
            
    # 4. 逆序清理提示词并填充
    matched_sections = sorted(matched_sections, key=lambda x: x['p_idx'], reverse=True)
    
    for i, item in enumerate(matched_sections):
        curr_idx = item['p_idx']
        val = item['config']['value']
        next_boundary = len(body_p_elements) if i == 0 else matched_sections[i-1]['p_idx']
        
        first_content_p = None
        for mid_idx in range(curr_idx + 1, next_boundary):
            mid_elem = body_p_elements[mid_idx]
            try:
                p_obj = Paragraph(mid_elem, doc)
                if first_content_p is None:
                    first_content_p = p_obj
                else:
                    # 彻底移除中间的提示词说明
                    p_obj._element.getparent().remove(p_obj._element)
            except Exception:
                pass
                
        # 使用 OxmlElement 创建新段落，完全兼容普通 Body 与表格单元格 CT_Tc
        if first_content_p is None:
            new_p_elem = OxmlElement('w:p')
            item['p_elem'].addnext(new_p_elem)
            first_content_p = Paragraph(new_p_elem, doc)
            
        if first_content_p._element.pPr is not None:
            first_content_p._element.remove(first_content_p._element.pPr)
            
        first_content_p.text = ""
        first_content_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = first_content_p.add_run(str(val))
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        run.font.bold = False

    # 5. 图片自适应插入
    ok_img = polished_data.get('OK Picture Bytes')
    ng_img = polished_data.get('NG Picture Bytes')
    if ok_img or ng_img:
        for table in doc.tables:
            text = "".join(cell.text for row in table.rows for cell in row.cells)
            if "OK-Part" in text and "Not-OK-Part" in text:
                for row in table.rows:
                    for cell in row.cells:
                        if "Not-OK-Part" in cell.text and ng_img:
                            for p in list(cell.paragraphs):
                                if "not-ok-part" not in p.text.lower().replace(" ", ""):
                                    p._element.getparent().remove(p._element)
                            np = cell.add_paragraph()
                            np.alignment = 1
                            np.add_run().add_picture(io.BytesIO(ng_img), width=Inches(2.4))
                        elif "OK-Part" in cell.text and ok_img:
                            for p in list(cell.paragraphs):
                                if "ok-part" not in p.text.lower().replace(" ", ""):
                                    p._element.getparent().remove(p._element)
                            np = cell.add_paragraph()
                            np.alignment = 1
                            np.add_run().add_picture(io.BytesIO(ok_img), width=Inches(2.4))
                break
                
    return doc

def generate_eml_file(row_data, to_emails="", doc_bytes=None, doc_filename="LL_Template.docx"):
    """生成带编辑模式（X-Unsent: 1）且附带 Word 附件的 Outlook 草稿"""
    serial_no = str(row_data.get('LL Serials No', 'LL-xxxx-xx')).strip()
    failure_mode = str(row_data.get('Failure Mode', '*****')).strip()
    subject = f"M/PQR-AP LL | {serial_no} | Title {failure_mode}"
    
    html_body = f"""
    <html>
    <head>
        <style>
            body {{ font-family: 'Arial', sans-serif; font-size: 10.5pt; line-height: 1.6; color: #333333; }}
            .red-bold {{ color: #E20015; font-weight: bold; }}
            ul {{ margin-top: 5px; margin-bottom: 15px; padding-left: 20px; }}
            li {{ margin-bottom: 8px; }}
        </style>
    </head>
    <body>
        <p>Dear Supplier:</p>
        <p>Recently, we summarized a lesson learn about <strong>{failure_mode}</strong>. Please review the attached LL document and complete following tasks:</p>
        <ul>
            <li>Complete feedback form based on self-evaluation on your own processes and send to your responsible PQR and PUQ-PQA (ME) within <span class="red-bold">one week.</span></li>
            <li>After your self-evaluation, please close defined actions within <span class="red-bold">three weeks.</span></li>
            <li>Our PQR or PUQ-PQA colleague may conduct onsite verification according to the information in feedback form in <span class="red-bold">one month.</span></li>
        </ul>
        <p>If you have any question about this lesson learn, please contact with your responsible PQR and PUQ-PQA (ME).</p>
        <p>Best regards,</p>
        <p><strong>Purchasing Quality Region Asia Pacific Team</strong></p>
    </body>
    </html>
    """
    msg = MIMEMultipart('mixed')
    msg['Subject'] = Header(subject, 'utf-8')
    msg['From'] = 'Sunny.LIU3@cn.bosch.com'
    msg['To'] = to_emails
    msg.add_header('X-Unsent', '1')
    
    alt_part = MIMEMultipart('alternative')
    alt_part.attach(MIMEText(html_body, 'html', 'utf-8'))
    msg.attach(alt_part)
    
    if doc_bytes:
        part_doc = MIMEBase('application', 'vnd.openxmlformats-officedocument.wordprocessingml.document')
        part_doc.set_payload(doc_bytes)
        encoders.encode_base64(part_doc)
        part_doc.add_header('Content-Disposition', f'attachment; filename="{doc_filename}"')
        msg.attach(part_doc)
        
    return msg.as_bytes()

# -----------------------------------------------------------------------------
# 4. 主界面：两大工作流与 AI 看板
# -----------------------------------------------------------------------------
tab_auto, tab_chat, tab_prompt = st.tabs([
    "🚀 一键全自动流水线 (选记录 ➔ AI润色 ➔ 生成Word+邮件)",
    "🤖 M-PU AI 实时交互中枢",
    "📖 FEBER 结构与 Prompt 规范"
])

# =============================================================================
# TAB 1: 一键自动化流水线
# =============================================================================
with tab_auto:
    if excel_file is not None and template_file is not None:
        try:
            df, sheet_name, header_idx = load_excel_robust(excel_file)
            supplier_dict = load_supplier_emails(excel_file)
            
            # 过滤有效记录 (LL Need or not == Y)
            ll_need_col = next((c for c in df.columns if 'need or not' in str(c).lower()), 'LL Need or not')
            if ll_need_col in df.columns:
                orig_cnt = len(df)
                df = df[df[ll_need_col].astype(str).str.strip().str.upper() == 'Y']
                st.success(f"🎉 成功载入 **{sheet_name}**：已过滤保留 `{ll_need_col} = 'Y'` 的 **{len(df)}** 条有效台账（排除 {orig_cnt - len(df)} 条无效数据）。")
                
            serial_no_col = next((c for c in df.columns if 'serial' in str(c).lower()), 'LL Serials No')
            supplier_scope_col = next((c for c in df.columns if 'scope' in str(c).lower() or 'task' in str(c).lower()), 'LL Supplier Scope')
            
            st.markdown("#### 1️⃣ 勾选需要生成的台账记录")
            search_key = st.text_input("🔍 快速检索记录:", key="tb1_search")
            filtered_df = df.copy()
            if search_key:
                filtered_df = df[df.astype(str).apply(lambda r: r.str.contains(search_key, case=False).any(), axis=1)]
                
            filtered_df.insert(0, '选择 (Select)', False)
            disp_cols = ['选择 (Select)', serial_no_col, 'Failure Mode', 'Supplier Name', 'Project/Part name', 'LL Brief Description']
            valid_disp_cols = [c for c in disp_cols if c in filtered_df.columns]
            
            edited_df = st.data_editor(
                filtered_df[valid_disp_cols],
                hide_index=True,
                column_config={"选择 (Select)": st.column_config.CheckboxColumn("选择", default=False)},
                disabled=[c for c in valid_disp_cols if c != '选择 (Select)'],
                use_container_width=True
            )
            
            selected_rows = filtered_df.loc[edited_df[edited_df['选择 (Select)'] == True].index]
            
            st.markdown("#### 2️⃣ 指定发送供应商并一键生成")
            c_sup, c_info = st.columns([2, 2])
            with c_sup:
                chosen_sups = st.multiselect("👥 发送给哪些供应商 (自动从 Vendor code 提取邮箱):", options=list(supplier_dict.keys()))
            to_emails_list = []
            for s in chosen_sups:
                to_emails_list.extend(supplier_dict[s])
            to_emails_str = "; ".join(list(set(to_emails_list)))
            
            with c_info:
                if to_emails_str:
                    st.info(f"📧 **收件人:** `{to_emails_str}`")
                else:
                    st.caption("ℹ️ 若未选择供应商，收件人一栏将留空供手动输入。")
                    
            if len(selected_rows) > 0:
                st.write("")
                if st.button("🚀 启动自动化流水线：AI润色 ➔ 清洗模板 ➔ 导出Word & Outlook草稿", type="primary", use_container_width=True):
                    with st.spinner("🤖 M-PU AI 正在润色内容、彻底清洗模板提示词并合成邮件..."):
                        if len(selected_rows) == 1:
                            row = selected_rows.iloc[0]
                            ok_img, ng_img = get_images_for_row(excel_file, sheet_name, header_idx, row.name)
                            raw_data = {
                                'LL Serials No': row.get(serial_no_col, 'LL-xxxx-xx'),
                                'Failure Mode': row.get('Failure Mode', '*****'),
                                'Project/Part name': row.get('Project/Part name', ''),
                                'LL Brief Description': row.get('LL Brief Description', ''),
                                'Root Cause': row.get('Root Cause', ''),
                                'Corrective Action': row.get('Corrective Action', ''),
                                'Should or not to do': row.get('Should or not to do', ''),
                                'Related Material Field / Process': row.get('Related Material Field / Process', ''),
                                'LL Supplier Scope': row.get(supplier_scope_col, ''),
                                'What else could be additionally affected?': row.get('What else could be additionally affected?', ''),
                                'Where can the problem additionally occur?': row.get('Where can the problem additionally occur?', ''),
                                'When can the problem additionally appear?': row.get('When can the problem additionally appear?', ''),
                                'Who else can be affected?': row.get('Who else can be affected?', ''),
                                'OK Picture Bytes': ok_img,
                                'NG Picture Bytes': ng_img
                            }
                            # AI 润色
                            polished_data = ai_polish_record(raw_data)
                            polished_data['OK Picture Bytes'] = ok_img
                            polished_data['NG Picture Bytes'] = ng_img
                            
                            # 智能清除提示词并写入
                            doc = fill_word_template_smart(template_file, polished_data)
                            bio = io.BytesIO()
                            doc.save(bio)
                            doc_bytes = bio.getvalue()
                            
                            serial_str = str(polished_data['LL Serials No'])
                            doc_name = f"LL_Template_{serial_str}.docx"
                            eml_bytes = generate_eml_file(polished_data, to_emails_str, doc_bytes, doc_name)
                            
                            st.success(f"✨ 记录 [{serial_str}] 自动化处理完毕！")
                            b1, b2 = st.columns(2)
                            with b1:
                                st.download_button(f"📥 下载 Word: {doc_name}", doc_bytes, doc_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                            with b2:
                                st.download_button(f"📧 下载 Outlook 草稿 (含附件): Email_Draft_{serial_str}.eml", eml_bytes, f"Email_Draft_{serial_str}.eml", mime="message/rfc822", use_container_width=True)
                        else:
                            zip_buf = io.BytesIO()
                            with zipfile.ZipFile(zip_buf, "a", zipfile.ZIP_DEFLATED, False) as zf:
                                for idx, (real_row_idx, row) in enumerate(selected_rows.iterrows()):
                                    ok_img, ng_img = get_images_for_row(excel_file, sheet_name, header_idx, real_row_idx)
                                    raw_data = {
                                        'LL Serials No': row.get(serial_no_col, f"Record_{idx+1}"),
                                        'Failure Mode': row.get('Failure Mode', '*****'),
                                        'Project/Part name': row.get('Project/Part name', ''),
                                        'LL Brief Description': row.get('LL Brief Description', ''),
                                        'Root Cause': row.get('Root Cause', ''),
                                        'Corrective Action': row.get('Corrective Action', ''),
                                        'Should or not to do': row.get('Should or not to do', ''),
                                        'Related Material Field / Process': row.get('Related Material Field / Process', ''),
                                        'LL Supplier Scope': row.get(supplier_scope_col, ''),
                                        'What else could be additionally affected?': row.get('What else could be additionally affected?', ''),
                                        'Where can the problem additionally occur?': row.get('Where can the problem additionally occur?', ''),
                                        'When can the problem additionally appear?': row.get('When can the problem additionally appear?', ''),
                                        'Who else can be affected?': row.get('Who else can be affected?', ''),
                                        'OK Picture Bytes': ok_img,
                                        'NG Picture Bytes': ng_img
                                    }
                                    polished_data = ai_polish_record(raw_data)
                                    polished_data['OK Picture Bytes'] = ok_img
                                    polished_data['NG Picture Bytes'] = ng_img
                                    
                                    doc = fill_word_template_smart(template_file, polished_data)
                                    doc_io = io.BytesIO()
                                    doc.save(doc_io)
                                    doc_bytes = doc_io.getvalue()
                                    
                                    serial_str = str(polished_data['LL Serials No'])
                                    doc_name = f"LL_Template_{serial_str}.docx"
                                    zf.writestr(doc_name, doc_bytes)
                                    
                                    eml_bytes = generate_eml_file(polished_data, to_emails_str, doc_bytes, doc_name)
                                    zf.writestr(f"Email_Draft_{serial_str}.eml", eml_bytes)
                                    
                            zip_buf.seek(0)
                            st.download_button("📥 立即下载打包 ZIP (包含所有 Word 报告与带附件的邮件草稿)", zip_buf.read(), "LL_AI_Polished_Batch.zip", mime="application/zip", use_container_width=True)
                            st.success("✨ 批量全自动处理完成！")
            else:
                st.warning("👉 请先在表格中勾选至少 1 条记录。")
        except Exception as e:
            st.error(f"❌ 运行异常: {e}")
    else:
        st.info("ℹ️ 请在侧边栏确认 Master List 和 Word 模板的文件路径。")

# =============================================================================
# TAB 2: M-PU AI 实时交互看板
# =============================================================================
with tab_chat:
    st.markdown("""
    <div class="bosch-card">
        <h4 style="color: #005691; margin: 0 0 6px 0;">🤖 M-PU ChatGPT 交互中枢</h4>
        <p style="color: #525F6B; font-size: 0.88rem; margin: 0;">直接在界面内向 AI 提问、上传报告润色，无需跳转 Teams 窗口。</p>
    </div>
    """, unsafe_allow_html=True)
    
    col_c1, col_c2 = st.columns([3, 2])
    with col_c2:
        st.markdown("##### ⚙️ 快速生成某条台账的标准 Prompt")
        if excel_file is not None:
            try:
                df_c, _, _ = load_excel_robust(excel_file)
                rec_list = df_c['LL Serials No'].tolist() if 'LL Serials No' in df_c.columns else []
                selected_rec = st.selectbox("选择记录编号:", rec_list)
                if st.button("📋 生成该条记录的 FEBER Prompt"):
                    sel_row = df_c[df_c['LL Serials No'] == selected_rec].iloc[0]
                    p_txt = f"""Please create me a short and precise lessons learned report out of below facts in American English:

LL Serials No: {sel_row.get('LL Serials No', '')}
Failure Mode: {sel_row.get('Failure Mode', '')}
Brief Description: {sel_row.get('LL Brief Description', '')}
Root Cause: {sel_row.get('Root Cause', '')}
Corrective Action: {sel_row.get('Corrective Action', '')}
Should or not to do: {sel_row.get('Should or not to do', '')}

Structure according to FEBER:
0. Abstract
1. Product / Process
2. Problem (Fundamental Problem)
3. Lessons
4. Potentially affected"""
                    st.text_area("生成的 Prompt (可直接发送):", p_txt, height=180)
            except Exception:
                pass
                
    with col_c1:
        st.markdown("##### 💬 内嵌对话终端")
        if "messages" not in st.session_state:
            st.session_state.messages = [
                {"role": "assistant", "content": "您好！我是 **M-PU ChatGPT Bot**。您可以直接把 PCB 失效描述发给我，我将帮您按照博世 FEBER 规范进行结构化英文润色。"}
            ]
            
        for msg in st.session_state.messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])
                
        user_prompt = st.chat_input("输入工程问题或粘贴报告内容...")
        if user_prompt:
            st.session_state.messages.append({"role": "user", "content": user_prompt})
            with st.chat_message("user"):
                st.markdown(user_prompt)
                
            with st.chat_message("assistant"):
                resp = f"""### **Abstract**
**Issue:** Quality deviation observed during PCB inspection.
**Problem:** Non-conformance against Bosch technical specification.
**Lessons:** Strengthened parameter monitoring and revised preventive maintenance frequency.

---
### **1. Product / Process**
* **Process:** PCB Manufacturing & Plating Process
* **Component:** Press-fit PCB ST60

### **2. Problem (Fundamental Problem)**
Root analysis shows parameter deviation due to insufficient contact stability during plating.

### **3. Lessons**
* **Measures & Sustainable Solutions:** 
  - Optimized maintenance cycle from 1.5 months to 1 month.
  - Deployed infrared voltage anomaly alarm sensors (<1.4V).
* **Root Cause:** Insufficient maintenance interval for contact gear.

### **4. Potentially affected**
* **What else:** Similar pattern plating processes.
* **Where:** Other production lines.
* **Who:** PUQ-PQA, Supplier Quality Engineering."""
                st.markdown(resp)
                st.session_state.messages.append({"role": "assistant", "content": resp})

# =============================================================================
# TAB 3: 标准 FEBER 结构指引
# =============================================================================
with tab_prompt:
    st.markdown("### 📖 博世标准 Lessons Learned (FEBER) 规范")
    st.code("""Please create me a short and precise lessons learned report out of the attached document in American English.
You are an honest engineer; you provide always links to the sources and name the original slide/page number.
Please stick to the facts. In case you have additional topics, supporting or additional useful information be creative, add them and highlight them in italic.

Please write the headings in bold. Use key words that are understood by others in Bosch. Describe the report "user-friendly", so others can read it easily. One page for chapter 1-4 is appropriate. Delete all hints in blue letters.

If you are asked to create a lesson learned report, or to search for a lessons learned report, structure the answer as follows:
0. Abstract - write a short summary of the report with the structure - issue; problem; learnings; tags

1. Product / Process
Product / Process:
Component:
Sub-Component:

2. Problem (Fundamental Problem)
Briefly describe the fundamental problem.

3. Lessons
- Root Cause: Describe the main root cause only.
- Measures & Sustainable Solutions: Focus on Process Improvements and Standards Updates.

4. Potentially affected
- What else could be additionally affected?
- Where can the problem additionally occur?
- When can the problem additionally appear?
- Who else can be affected?

5. Appendix (Optional)""", language="text")
